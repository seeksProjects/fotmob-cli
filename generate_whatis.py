"""Generate whatis.docx — comprehensive project state assessment."""
from docx import Document
from docx.shared import Pt
import datetime

doc = Document()
doc.add_heading("FotMob CLI - Project State Assessment", level=0)
doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
doc.add_paragraph("Comprehensive audit of current state: what works, what is broken, what is difficult.")
doc.add_paragraph("")

# ====== 1. PROJECT OVERVIEW ======
doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "The FotMob CLI is a football data tool that queries FotMob internal API, scrapes match pages "
    "via headless Chrome, and uses AI (Groq/Llama 3.3 70B) to provide natural language answers. "
    "Runs as: terminal CLI, Flask web app, deployed on Render.com (Docker)."
)
doc.add_paragraph("Files: 14 Python files, ~5000 lines total.")

t = doc.add_table(rows=15, cols=4)
t.style = "Table Grid"
for i, h in enumerate(["File", "Lines", "Purpose", "Status"]):
    t.rows[0].cells[i].text = h
data = [
    ("api.py", "108", "FotMob API client with caching", "Working"),
    ("cli.py", "488", "Click CLI command handlers", "Working"),
    ("config.py", "74", "Config (env vars, .env, config.json)", "Working"),
    ("display.py", "622", "Rich terminal output formatting", "Working"),
    ("browser.py", "199", "Seleniumbase Chrome for Turnstile bypass", "Working (slow)"),
    ("nlp.py", "590", "NLP intent detection + fuzzy matching", "Working"),
    ("interactive.py", "1063", "Interactive REPL with LLM + keyword NLP", "Working"),
    ("llm_parser.py", "237", "LLM parser (Gemini/Groq fallback chain)", "Working"),
    ("ai_answer.py", "372", "Two-pass AI answer generation", "Working"),
    ("live.py", "231", "Live match tracking with polling", "Working"),
    ("match_data.py", "400+", "Comprehensive match data extraction", "Working locally, issues on cloud"),
    ("odds.py", "146", "Betting odds via The Odds API", "Working"),
    ("tv_channels.py", "204", "TV channels via LiveSoccerTV scraping", "Working locally, fragile on cloud"),
    ("web.py", "770", "Flask web server + dual UI (chat/terminal)", "BROKEN on cloud"),
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        t.rows[i + 1].cells[j].text = val

# ====== 2. WHAT WORKS ======
doc.add_heading("2. What Works (Confirmed)", level=1)
works = {
    "Terminal CLI (python cli.py)": "Interactive mode with natural language queries. All commands work: search, standings, fixtures, team, squad, player, match, transfers, news, top-scorers, live tracking, commentary.",
    "AI-Curated Answers (Terminal)": "Two-pass system: Groq parses intent, fetches FotMob data, Groq curates a specific conversational answer. Handles typos, complex queries, ambiguous phrasing.",
    "Turnstile Bypass": "Seleniumbase headless Chrome loads FotMob pages and extracts __NEXT_DATA__ JSON. Gets data the API blocks (matchDetails, playerData).",
    "Comprehensive Match Data (Terminal)": "Extracts: lineups with shirt numbers, coaches, referee, stadium, weather, injuries, insights, H2H, team form, top scorers, events, live stats, squad market values, TV channels, betting odds.",
    "Live Commentary (Terminal)": "Scrapes Commentary tab from match pages. Returns minute-by-minute text. AI summarizes conversationally.",
    "Betting Odds": "Real odds from The Odds API (free, 500 req/month). 7 major leagues, 3 bookmakers per match.",
    "TV Channels (Terminal/Local)": "LiveSoccerTV scraping works locally. Gets US/Europe/Americas channels.",
    "Fuzzy Matching + NLP": "200+ phrases, thefuzz (threshold 70), league/team dicts with misspellings.",
    "Web UI Chat Mode": "WhatsApp-style bubbles, timestamps, avatar. Responsive on mobile.",
    "Cloud Deployment": "Docker on Render.com free tier. Auto-deploys from GitHub.",
}
for title, desc in works.items():
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ====== 3. WHAT IS BROKEN ======
doc.add_heading("3. What Is Broken", level=1)
broken = {
    "Web handler 'vs' queries fail on cloud": (
        "PROBLEM: 'Villarreal vs Real Sociedad watch' returns 'Could not generate answer' on web "
        "but works perfectly on terminal CLI.\n\n"
        "ROOT CAUSE: web.py and interactive.py have DUPLICATED query handling logic that has diverged. "
        "The terminal version (interactive.py) has been updated more thoroughly with match_data.py "
        "integration, TV channels, odds, commentary. The web version (web.py) was updated separately "
        "but the browser on Render cloud is slower and may timeout. The web handler also uses older "
        "event extraction code in some paths instead of the comprehensive match_data.py extractor."
    ),
    "Terminal vs Web response inconsistency": (
        "Terminal gives full detailed responses (referee, TV, odds, lineups, injuries). "
        "Web gives minimal or error responses for the same queries. "
        "Root cause: two separate codebases doing the same thing differently."
    ),
    "Terminal UI in web toggle looks broken": (
        "The terminal mode CSS overrides conflict with chat mode message structure. "
        "Switching from chat to terminal doesn't transform existing bubbles cleanly."
    ),
    "TV channels on cloud (Render)": (
        "LiveSoccerTV scraping requires browser. Render 512MB RAM + Chrome for FotMob + "
        "Chrome for LiveSoccerTV is unreliable. TV lookup often fails silently."
    ),
    "Prediction template placeholders unfilled": (
        "Poll text shows: 'Gremio have VitoriaW, {2}D, 1L...' — template filling logic "
        "in match_data.py does not handle all parameter positions."
    ),
}
for title, desc in broken.items():
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ====== 4. WHAT IS DIFFICULT ======
doc.add_heading("4. What Is Difficult / Cannot Be Done", level=1)
difficult = {
    "Browser performance on Render free tier": (
        "Chrome needs ~300MB. Render free tier = 512MB total. First browser launch = 10-30s. "
        "Each page load = 4-6s. Chrome + Python + Flask on 512MB is tight. "
        "This affects: match details, player stats, TV channels, commentary."
    ),
    "FotMob API data is stale for live matches": (
        "Team/league API endpoints return DELAYED scores for ongoing matches, especially lower leagues. "
        "Only browser-scraped match page has real-time data. Score queries can be wrong without browser."
    ),
    "TV channel data is location-dependent": (
        "FotMob truncates TV names server-side. LiveSoccerTV returns channels based on server location "
        "(Render = Oregon, US). Users in Africa/Asia get US channels. No way to set user country without input."
    ),
    "Betting odds limited to major leagues": (
        "The Odds API free tier covers EPL, La Liga, Bundesliga, Serie A, Ligue 1, UCL, UEL. "
        "Turkish, Brazilian, Indian leagues return nothing."
    ),
    "FotMob can change API anytime": (
        "All endpoints are internal/undocumented. URL patterns, JSON structure, anti-bot measures "
        "can change without notice. The Turnstile bypass and data extraction are reverse-engineered."
    ),
    "Heatmaps, shot maps, momentum graphs": (
        "Data is available but CANNOT be rendered in terminal or text. These are visual-only features. "
        "Would need a web frontend with canvas/SVG to display them."
    ),
    "No user authentication on web app": (
        "Anyone with the URL can use it. API quotas (Groq 14,400/day, Odds API 500/month) "
        "will be exhausted if the app gets traffic."
    ),
    "Exact betting odds from FotMob itself": (
        "1xBET widget on FotMob specifically blocks headless browsers. "
        "The odds visible on the website are NOT in the API or rendered DOM. "
        "We use The Odds API as alternative."
    ),
}
for title, desc in difficult.items():
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ====== 5. CODE QUALITY ISSUES ======
doc.add_heading("5. Code Quality Issues", level=1)
quality = {
    "CRITICAL: Duplicated code between web.py and interactive.py": (
        "Both files have their own: _resolve_team(), _resolve_player(), query routing, "
        "'vs' handler, match preview handler, commentary handler. "
        "Bug fixed in one is NOT fixed in the other. This is the #1 source of "
        "'works on terminal, breaks on web' bugs."
    ),
    "Hardcoded league IDs in 5 different places": (
        "cli.py: 7 leagues. interactive.py: 21 leagues. web.py: 7 and 13 in different functions. "
        "Adding a new league requires changes in 5 places."
    ),
    "4 different league/team mapping dictionaries": (
        "nlp.py LEAGUE_DICT (47 entries), tv_channels.py TEAM_SLUGS (47 teams), "
        "odds.py LEAGUE_TO_SPORT (12 entries), cli.py LEAGUE_IDS (11 entries). "
        "Overlapping, inconsistent, maintained separately."
    ),
    "NLP uses substring matching, not word boundaries": (
        "'live' matches inside 'delivery', 'rival'. Should use regex word boundaries."
    ),
    "Unbounded API cache": (
        "api.py cache dict grows forever. No max size or LRU eviction. "
        "Memory leak risk on long-running web server."
    ),
    "Global state mutation for cache control": (
        "live.py and interactive.py directly mutate api._cache and api._cache_ttl. "
        "Thread-unsafe, side effects between functions."
    ),
    "No rate limit handling or retry logic": (
        "No exponential backoff for FotMob API, Groq API, or Odds API. "
        "429 errors cause immediate failure instead of retry."
    ),
}
for title, desc in quality.items():
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ====== 6. FEATURE COVERAGE ======
doc.add_heading("6. Feature Coverage vs FotMob Website", level=1)
t2 = doc.add_table(rows=22, cols=3)
t2.style = "Table Grid"
t2.rows[0].cells[0].text = "Feature"
t2.rows[0].cells[1].text = "Status"
t2.rows[0].cells[2].text = "Notes"
features = [
    ("Live Scores", "WORKING", "Via API + browser for real-time"),
    ("Live Commentary", "WORKING", "Scraped from Commentary tab"),
    ("Match Events", "WORKING", "Goals, cards, subs via browser"),
    ("xG / Expected Goals", "WORKING", "Data extracted, shown in stats"),
    ("Player Ratings", "PARTIAL", "In lineups, not always populated"),
    ("Heatmaps", "IMPOSSIBLE", "Visual only, cannot render in text"),
    ("Momentum Graph", "IMPOSSIBLE", "Visual only, data available"),
    ("League Standings", "WORKING", "Full 20-team tables"),
    ("Team Form", "WORKING", "Last 5 matches with results"),
    ("Team Squad", "WORKING", "Full roster with shirt numbers"),
    ("Player Stats", "WORKING", "Season stats + recent matches via browser"),
    ("Predicted Lineups", "WORKING", "Formation + players + shirt numbers"),
    ("Injuries/Suspended", "WORKING", "Full list with return dates"),
    ("Referee", "WORKING", "Name and country"),
    ("Stadium/Weather", "WORKING", "Full venue info + forecast"),
    ("Coaches", "WORKING", "Both coaches"),
    ("H2H Record", "WORKING", "Summary + recent meetings"),
    ("TV Channels", "PARTIAL", "Works locally, fragile on cloud"),
    ("Betting Odds", "WORKING", "Via The Odds API (major leagues)"),
    ("Transfers", "WORKING", "Global + team-specific"),
    ("News", "WORKING", "World football headlines"),
]
for i, (f, s, n) in enumerate(features):
    for j, v in enumerate([f, s, n]):
        t2.rows[i + 1].cells[j].text = v

# ====== 7. EXTERNAL DEPENDENCIES ======
doc.add_heading("7. External Dependencies & API Keys", level=1)
t3 = doc.add_table(rows=6, cols=4)
t3.style = "Table Grid"
for i, h in enumerate(["Service", "Free Tier", "Usage", "Risk"]):
    t3.rows[0].cells[i].text = h
deps = [
    ("Groq (Llama 3.3 70B)", "14,400 req/day", "Query parsing + answer (2 calls/query)", "Low"),
    ("Gemini", "~1,500 req/day", "Backup LLM (currently exhausted)", "Medium"),
    ("The Odds API", "500 req/month", "Betting odds", "Medium"),
    ("FotMob", "Unlimited (unofficial)", "All football data", "HIGH - can block anytime"),
    ("LiveSoccerTV", "Unlimited (scraping)", "TV channels", "HIGH - Cloudflare, fragile"),
]
for i, row_data in enumerate(deps):
    for j, v in enumerate(row_data):
        t3.rows[i + 1].cells[j].text = v

# ====== 8. PRIORITY RECOMMENDATIONS ======
doc.add_heading("8. Priority Recommendations", level=1)
doc.add_heading("P1 (Fix First):", level=2)
doc.add_paragraph(
    "1. UNIFY web.py and interactive.py query handling into a single module (query_handler.py). "
    "Both should call the same functions. This eliminates ALL 'works on terminal, breaks on web' bugs.\n\n"
    "2. Fix the web handler to support all actions: upcoming, commentary, match_preview with full match_data.py.\n\n"
    "3. Fix prediction template placeholder filling in match_data.py."
)
doc.add_heading("P2 (Refactor):", level=2)
doc.add_paragraph(
    "4. Centralize all league IDs, team slugs, and mappings into one constants.py file.\n\n"
    "5. Add cache size limit (LRU, max 500 entries) to api.py.\n\n"
    "6. Add word boundary matching to NLP intent detection.\n\n"
    "7. Add proper error messages when cloud features fail (TV, browser, etc)."
)
doc.add_heading("P3 (Polish):", level=2)
doc.add_paragraph(
    "8. Add rate limit retry logic for Groq and Odds API.\n\n"
    "9. Add user country selection for TV channels.\n\n"
    "10. Document config.json format and setup instructions in README."
)

# Save
path = "C:/Users/USER/Desktop/assignments/fotmob_in_terminal/whatis.docx"
doc.save(path)
print(f"Saved to {path}")
